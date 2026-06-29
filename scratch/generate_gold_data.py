import os
import json

# Ensure directories exist
os.makedirs("tests/gold", exist_ok=True)

# 1. recruiter.csv
csv_content = """name,email,phone,current_company,title,skills
John Doe,john@gmail.com,9876543210,Google,Software Engineer,Java;Spring Boot"""
with open("tests/gold/recruiter.csv", "w", encoding="utf-8") as f:
    f.write(csv_content)

# 2. ats.json
ats_content = {
    "candidateName": "John A Doe",
    "contactEmail": "john@gmail.com",
    "currentEmployer": "Google LLC",
    "jobTitle": "Software Engineer II",
    "skills": ["Java", "AWS"]
}
with open("tests/gold/ats.json", "w", encoding="utf-8") as f:
    json.dump(ats_content, f, indent=2)

# 3. notes.txt
notes_content = """Candidate Name: John Doe
Spoke with candidate. Currently at Google.
Strong skills in Java, spring boot, AWS.
Email is john@gmail.com
Phone is 9876543210"""
with open("tests/gold/notes.txt", "w", encoding="utf-8") as f:
    f.write(notes_content)

# 4. Generate resume.docx using python-docx
try:
    import docx
    doc = docx.Document()
    doc.add_paragraph("John Andrew Doe")
    doc.add_paragraph("Email: john@gmail.com")
    doc.add_paragraph("Phone: (+91)-98765-43210")
    doc.add_paragraph("Headline: Senior Software Engineer")
    doc.add_paragraph("Experience:")
    doc.add_paragraph("Worked at Google Inc. as Software Engineer.")
    doc.add_paragraph("Skills:")
    doc.add_paragraph("Java")
    doc.add_paragraph("Spring Boot")
    doc.add_paragraph("AWS")
    doc.save("tests/gold/resume.docx")
    print("Generated docx successfully.")
except Exception as e:
    # Fallback to simple txt or copy mock file
    print(f"Error generating docx: {e}")
